from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_h(doc, text, level=1, color='1a2e4a'):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_p(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def header_row(table, headers, bg='1a2e4a'):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], bg)
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                run.bold = True

def make_table(doc, headers, rows, bg='1a2e4a'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    header_row(t, headers, bg)
    for row_data in rows:
        r = t.add_row().cells
        for i, val in enumerate(row_data):
            r[i].text = val
    return t

def style_doc(doc):
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    s = doc.sections[0]
    s.page_height = Cm(29.7)
    s.page_width = Cm(21)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)

# =========================================================
# Cahier des Charges
# =========================================================
doc = Document()
style_doc(doc)

today = datetime.date.today().strftime("%d %B %Y")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CAHIER DES CHARGES TECHNICO-FONCTIONNEL')
run.bold = True; run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1a, 0x2e, 0x4a)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('PROJET SENTINELLE — MedSPAD 2026')
run2.bold = True; run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x0d, 0x6e, 0xfd)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Plateforme d\'Intelligence Epidemiologique et de Prevention des Addictions')
run3.italic = True; run3.font.size = Pt(13)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run('Version 1.0  |  ' + today).font.size = Pt(11)

doc.add_page_break()

# ------ SECTION 1 ------
add_h(doc, '1. Présentation Générale', 1)
add_p(doc, '1.1 Contexte', bold=True)
add_p(doc, 'Le projet SENTINELLE s\'inscrit dans le cadre de l\'enquête épidémiologique MedSPAD 2026 (Mediterranean School Survey Project on Alcohol and Other Drugs). Face à la complexité croissante des conduites addictives chez les jeunes en Tunisie (polyconsommation, nouvelles substances psychoactives, cyberaddiction) et à la nécessité d\'une réactivité accrue des pouvoirs publics, le projet vise à numériser l\'intégralité du cycle de vie des données, de la collecte par questionnaire papier jusqu\'à la prise de décision via des tableaux de bord analytiques.')

doc.add_paragraph()
add_p(doc, '1.2 Objectif du Document', bold=True)
add_p(doc, 'Ce cahier des charges définit les spécifications fonctionnelles, techniques et architecturales pour la conception, le développement et le déploiement de la plateforme web Sentinelle, qui servira de socle principal de gestion, d\'analyse et de visualisation de l\'enquête nationale MedSPAD 2026 pour les 24 gouvernorats.')

# ------ SECTION 2 ------
doc.add_paragraph()
add_h(doc, '2. Spécifications Fonctionnelles', 1)
add_p(doc, 'Le système est articulé autour de quatre modules métier principaux :')

doc.add_paragraph()
add_p(doc, '2.1 Module 1 : Pipeline de Numérisation et de Collecte (OCR)', bold=True)
mod1 = [
    'Remplacement de la saisie manuelle par un système de Reconnaissance Optique de Caractères (OCR) piloté par Tesseract.',
    'Formulaire interactif QuestionnaireForm (React) permettant une collecte 1:1 avec le format papier MedSPAD (21 sections, choix des échelles standardisées).',
    'Interface de Révision : Intégration d\'une boucle de validation humaine (ScanPage / IntegrityLab) permettant de corriger les anomalies avant soumission en base.',
    'Création et gestion des rapports de classe (ClassReport) pour encadrer administrativement chaque lot de questionnaires.'
]
for item in mod1:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(11)

doc.add_paragraph()
add_p(doc, '2.2 Module 2 : Moteur de Prédiction et IA', bold=True)
mod2 = [
    'Scoring Clinique : Déploiement d\'un pipeline Machine Learning (Voting Regressor via Scikit-Learn) prédisant les niveaux de stress et de vulnérabilité.',
    'Moteur de corrélation asynchrone (Celery) : extraction de co-occurrences (Apriori), matrices de sévérité (Chi-carré), profils démographiques.',
    'Clustering spatial (DBSCAN) : Détection de "hotspots" épidémiologiques sans divulguer les adresses exactes.'
]
for item in mod2:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(11)

doc.add_paragraph()
add_p(doc, '2.3 Module 3 : Visualisation (Dashboarding) et Cartographie', bold=True)
mod3 = [
    'Navigation radiale : Sélection des 21 sections du MedSPAD via une roue interactive centrale (RadialSectionWheel).',
    'Tableaux de bords (Homepage) : Synthèse globale incluant prévalence, âge moyen, stress moyen et section dominante.',
    'Cartographie (Map3D / TunisiaChoropleth) : Visualisation par gouvernorat (choroplèthe) ou par clusters (pins géographiques masqués).',
    'Laboratoires d\'intelligence : SocialLab, IntegrityLab, ComorbidityLab, RankingsLab pour les analyses avancées (Super Admin).'
]
for item in mod3:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(11)

doc.add_paragraph()
add_p(doc, '2.4 Module 4 : Gestion des Rôles (RBAC)', bold=True)
mod4 = [
    ('SUPER_ADMIN', 'Vision nationale complète, accès à tous les laboratoires (Labs), export brut, gestion des utilisateurs, export SIDRA.'),
    ('GLOBAL_ADMIN', 'Analyse nationale complète, lecture seule des données.'),
    ('REGIONAL_ANALYST', 'Filtrage strict sur le gouvernorat assigné. Accès aux analyses comparées (moyennes nationales anonymes).'),
    ('PRACTITIONER', 'Accès établissement. Gestion des soumissions, génération de QR codes, vue agrégée locale.'),
    ('OPERATOR', 'Saisie de données (Scan/OCR) uniquement.')
]
make_table(doc, ['Rôle', 'Droits & Périmètre'], mod4)

# ------ SECTION 3 ------
doc.add_paragraph()
add_h(doc, '3. Spécifications Techniques', 1)
add_p(doc, '3.1 Architecture du Système', bold=True)
stack = [
    ('Frontend', 'React 19 (Vite), Tailwind CSS, Framer Motion, Recharts, MapLibre GL'),
    ('Backend', 'Django 6.0, Django REST Framework 3.17'),
    ('Base de données', 'SQLite (dev) / PostgreSQL (production recommandée)'),
    ('Authentification', 'JWT (SimpleJWT) avec blocage de sécurité (failed_attempts) et journalisation AuditLog'),
    ('Data Science', 'Scikit-Learn, Pandas, Mlxtend, SciPy'),
    ('Traitement Asynchrone', 'Celery + Redis pour l\'exécution des algorithmes de corrélation lourds')
]
make_table(doc, ['Composant', 'Technologies'], stack)

doc.add_paragraph()
add_p(doc, '3.2 Exigences Non-Fonctionnelles', bold=True)
nfr = [
    'Performances : Le Dashboard (Homepage API) doit se charger en moins de 2 secondes. Les tâches d\'IA complexes (corrélations) utilisent un polling frontend toutes les 2 secondes.',
    'Sécurité / RGPD : Interdiction stricte de stocker des noms d\'élèves ou des identifiants directs. Les coordonnées GPS sont cryptées et rendues floues via l\'algorithme DBSCAN (centroïdes).',
    'Extensibilité : Le système "DynamicQuestion" permet d\'ajouter de nouvelles questions au formulaire sans modifier le code source.',
    'Interopérabilité : Un endpoint dédié (/api/sidra/export/) permet l\'échange de statistiques macro avec la plateforme nationale SIDRA.'
]
for item in nfr:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(11)

# ------ SECTION 4 ------
doc.add_paragraph()
add_h(doc, '4. Modèle de Données', 1)
add_p(doc, 'Le modèle de base est centré autour de l\'entité QuestionnaireSession, liée à un SchoolEstablishment, un Governorate, et un ClassReport. À chaque session correspond 21 modèles 1:1 (SectionA à SectionZ) reproduisant la structure de la grille MedSPAD.')
add_p(doc, 'Champs analytiques pré-calculés sur la session (booléens) : tobacco_user, ecig_user, hookah_user, alcohol_user, cannabis_user, etc., pour optimiser les temps de réponse de l\'API.')

# ------ SECTION 5 ------
doc.add_paragraph()
add_h(doc, '5. Livrables Attendus', 1)
livrables = [
    ('Code Source', 'Dépôts Frontend (React) et Backend (Django) avec requirements.txt et package.json à jour.'),
    ('Base de données', 'Scripts de migration Django et fixtures initiales (gouvernorats, typologie de questions).'),
    ('Modèles IA', 'Pipelines Scikit-Learn sérialisés et scripts d\'entraînement associés.'),
    ('Documentation', 'Le présent Cahier des Charges, le Project Initiation Document (PID), et un README technique d\'installation.')
]
make_table(doc, ['Livrable', 'Description'], livrables)

doc.save('Cahier_des_Charges_Sentinelle.docx')
print('CDC saved successfully.')
