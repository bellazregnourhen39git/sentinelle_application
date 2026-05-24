from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_h(doc, text, level=1, color='1a2e4a'):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_p(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def header_row(table, headers, bg='1a2e4a'):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], bg)
        for par in hdr[i].paragraphs:
            for run in par.runs:
                run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                run.bold = True

def make_table(doc, headers, rows, bg='1a2e4a'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    header_row(t, headers, bg)
    for row_data in rows:
        r = t.add_row().cells
        for i, val in enumerate(row_data):
            r[i].text = val
    return t

def style_doc(doc):
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    s = doc.sections[0]
    s.page_height = Cm(29.7)
    s.page_width = Cm(21)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)

# =========================================================
# PID
# =========================================================
doc = Document()
style_doc(doc)

today = datetime.date.today().strftime("%d %B %Y")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROJECT INITIATION DOCUMENT (PID)')
run.bold = True; run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1a, 0x2e, 0x4a)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('SENTINELLE — Plateforme MedSPAD 2026')
run2.bold = True; run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x0d, 0x6e, 0xfd)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Epidemiological Intelligence & Addiction Prevention Platform')
run3.italic = True; run3.font.size = Pt(13)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run('Version 1.0  |  ' + today).font.size = Pt(11)

doc.add_page_break()

# ------ SECTION 1 ------
add_h(doc, '1. Executive Summary', 1)
add_p(doc, 'SENTINELLE is a state-of-the-art epidemiological intelligence platform built for the MedSPAD 2026 survey across all 24 Tunisian governorates. It replaces slow, error-prone paper-based data workflows with a real-time AI-augmented decision-support system for public health authorities, addiction specialists, and education policymakers.')
doc.add_paragraph()
add_p(doc, 'The platform delivers four integrated capabilities: (1) an OCR-powered data ingestion pipeline with human-in-the-loop validation; (2) a Scikit-Learn Voting Regressor engine for student stress and pre-addiction risk scoring; (3) a multi-scope interactive dashboard with a 21-section radial navigation wheel and choropleth map; and (4) an asynchronous statistical correlation engine (Apriori, chi-square, DBSCAN) for deep behavioral pattern discovery.')
doc.add_paragraph()
add_p(doc, 'Built on Django 6.0 REST Framework and React 19 (Vite), the platform enforces a five-tier Role-Based Access Control architecture (SUPER_ADMIN, GLOBAL_ADMIN, REGIONAL_ANALYST, PRACTITIONER, OPERATOR) that ensures strict data compartmentalization and zero PII exposure for minors.')

# ------ SECTION 2 ------
doc.add_paragraph()
add_h(doc, '2. Project Context & Business Case', 1)
add_p(doc, '2.1  Problem Statement', bold=True)
add_p(doc, 'The MedSPAD survey programme has historically suffered from: (a) manual data entry bottlenecks causing weeks of delay before analysis is possible; (b) estimated 15-20% error rates in hand-transcribed questionnaires; (c) static retrospective reporting that makes public health responses reactive rather than preventive; and (d) zero cross-variable behavioral correlation capability across the 21 survey dimensions.')
doc.add_paragraph()
add_p(doc, '2.2  Strategic Opportunity', bold=True)
add_p(doc, 'The accelerating prevalence of New Psychoactive Substances (NPS), e-cigarettes, and polysubstance use among Tunisian youth, combined with scientifically documented links between academic stress and addiction onset, creates an urgent mandate for real-time predictive intelligence. SENTINELLE closes the gap between field-level paper collection and national policy intelligence by digitizing, correlating, and predicting risk in near real-time.')

# ------ SECTION 3 ------
doc.add_paragraph()
add_h(doc, '3. Project Objectives', 1)
objectives = [
    ('OBJ-01', 'Automate Data Ingestion', 'Eliminate manual entry via Tesseract OCR with strict JSON schema validation achieving 100% structural parity with the 21-section MedSPAD 2026 questionnaire.'),
    ('OBJ-02', 'Real-Time Dashboard', 'Deploy role-scoped interactive dashboard (5 user tiers) with RadialSectionWheel, choropleth maps, and KPI analytics covering all 24 governorates.'),
    ('OBJ-03', 'Predictive AI Engine', 'Implement Voting Regressor pipeline scoring student stress from socioeconomic, academic, and behavioral indicators to identify pre-addiction vulnerability.'),
    ('OBJ-04', 'Correlation Analytics', 'Provide Celery-backed async correlation engine with 4 modules: Apriori co-occurrence, chi-square severity matrix, demographic patterns, DBSCAN geospatial hotspots.'),
    ('OBJ-05', 'Data Privacy & Compliance', 'Zero-PII architecture: GPS data encrypted at rest, location outputs anonymized via DBSCAN cluster centroids, strict RBAC preventing cross-regional data leakage.'),
    ('OBJ-06', 'SIDRA Interoperability', 'Enable standardized national data export to SIDRA (Systeme d\'Information sur les Drogues et les Risques a l\'Adolescence) via authenticated API endpoint.'),
]
make_table(doc, ['ID', 'Objective', 'Detail'], objectives)

# ------ SECTION 4 ------
doc.add_paragraph()
add_h(doc, '4. Project Scope', 1)
add_p(doc, '4.1  In Scope', bold=True)
in_scope = [
    'Full 21-section MedSPAD 2026 questionnaire digitization (Sections A, B, C, D, E, G, H, I, J, K, L, M, N, P, Q, R, S, T, U, V, Z)',
    'OCR pipeline: Tesseract-based scan ingestion + human validation interface (ScanPage, IntegrityLab)',
    'Bilingual QuestionnaireForm (FR/AR) with all official choice scales from the paper form',
    'ClassReport administrative workflow: session tethering, finalization, QR code distribution',
    'Five-role RBAC: SUPER_ADMIN, GLOBAL_ADMIN, REGIONAL_ANALYST, PRACTITIONER, OPERATOR',
    'Sentinelle Dashboard: RadialSectionWheel, SectionDetailPanel, TunisiaChoropleth, national/regional/school scopes',
    'Intelligence Labs: SocialLab, IntegrityLab, ComorbidityLab, RankingsLab',
    'Analytics engine: 4-module correlation system (Co-occurrence, Severity Matrix, Demographics, DBSCAN)',
    'AI Stress Prediction: Voting Regressor model integration',
    '30+ REST API endpoints covering auth, questionnaire, stats, SIDRA, class reports, dynamic questions',
    'User Management: invite-token workflow, approval lifecycle, CSV export, AuditLog',
    'Dynamic Question Engine: admin-configurable questionnaire extension without code changes',
    'Platform Terminology: bilingual configurable labels (FR/AR) for Super Admin',
]
for item in in_scope:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

doc.add_paragraph()
add_p(doc, '4.2  Out of Scope', bold=True)
out_scope = [
    'Direct clinical intervention or individual student follow-up communications',
    'Collection or processing of direct personal identifiers (student names, national IDs, exact GPS coordinates)',
    'Real-time classroom administration (paper remains primary collection medium)',
    'Native mobile applications (iOS/Android)',
]
for item in out_scope:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

# ------ SECTION 5 ------
doc.add_paragraph()
add_h(doc, '5. Technical Architecture', 1)
add_p(doc, '5.1  Technology Stack', bold=True)
stack = [
    ('Backend', 'Django 6.0.4 + Django REST Framework 3.17.1 (Python)'),
    ('Authentication', 'JWT via djangorestframework-simplejwt 5.5.1 with SecureTokenObtainPairSerializer, brute-force lockout, and full AuditLog recording'),
    ('Database', 'SQLite (development); PostgreSQL-ready schema for production'),
    ('AI / ML', 'Scikit-Learn Voting Regressor pipeline, Pandas, Mlxtend (Apriori), SciPy (chi-square), DBSCAN'),
    ('OCR', 'Tesseract 5.x via pytesseract with custom pre-processing for MedSPAD form layouts'),
    ('Frontend', 'React 19 (Vite), Tailwind CSS, Framer Motion'),
    ('Data Visualisation', 'Recharts, D3.js, MapLibre GL (3D map), custom TunisiaChoropleth SVG'),
    ('Async Processing', 'Celery + Redis for correlation engine background tasks'),
    ('i18n', 'react-i18next (French / Arabic bilingual)'),
    ('CORS', 'django-cors-headers 4.9.0'),
]
make_table(doc, ['Component', 'Technology'], stack)

doc.add_paragraph()
add_p(doc, '5.2  Data Model Architecture', bold=True)
add_p(doc, 'The database schema mirrors the MedSPAD 2026 questionnaire with strict 1:1 field parity. The central entity is QuestionnaireSession, carrying computed boolean risk flags (tobacco_user, alcohol_user, cannabis_user, cocaine_user, ecstasy_user, heroin_user, inhalant_user, hookah_user, ecig_user, tranquilizer_user, has_risk_behavior) for O(1) dashboard aggregation. Each of the 21 survey sections has a dedicated OneToOneField-linked model (SectionA through SectionZ). Seven standardized choice-scale constants (FREQUENCY_LIFETIME, FREQUENCY_30DAYS_CIGS, AGE_SCALE, STRESS_FREQ, HONESTY_SCALE, YES_NO, SATISFACTION_SCALE) enforce data integrity system-wide. Supporting entities include: Governorate (24 wilayas), SchoolEstablishment, SchoolClass, ClassReport (administrative session metadata), DynamicQuestion (extensible questionnaire engine), AuditLog, and PlatformTerminology.')

doc.add_paragraph()
add_p(doc, '5.3  Application Pages & Components', bold=True)
pages = [
    ('LandingPage', 'Public marketing/entry page with role-based redirect after authentication'),
    ('Login', 'JWT authentication with secure token storage and role-based redirect'),
    ('SetPassword', 'Token-activated password setup for invited users'),
    ('PractitionerGuide', 'Operator/Practitioner landing hub: class report creation, QR distribution, scan/OCR, session collection'),
    ('ClassReportPage', 'Administrative class report creation/management with session tethering'),
    ('SessionCollectionView', 'Real-time questionnaire collection progress tracker per class report'),
    ('QRCodePage', 'QR code generation for mobile questionnaire distribution'),
    ('ScanPage', 'OCR scanning interface with Tesseract integration and human validation loop'),
    ('SentinelleDashboard', 'Main analytics hub: RadialSectionWheel + SectionDetailPanel + TunisiaMap + RegionalProfilePanel'),
    ('RegionalDeepDivePage', 'Governorate-level deep-dive with school cluster map and benchmarks'),
    ('SocialLab', 'Social behavior & digital addiction correlation analysis (SUPER_ADMIN/GLOBAL_ADMIN)'),
    ('IntegrityLab', 'Honesty index validation and response reliability analysis'),
    ('ComorbidityLab', 'Cross-substance comorbidity pattern discovery'),
    ('RankingsLab', 'National governorate rankings by substance prevalence'),
    ('UserManagement', 'Full user lifecycle management with invite, approve, reject, delete, CSV export'),
    ('SubmissionsViewer', 'Raw submission browser with per-submission section detail view'),
]
make_table(doc, ['Page/Component', 'Function'], pages)

# ------ SECTION 6 ------
doc.add_paragraph()
add_h(doc, '6. RBAC Security Architecture', 1)
roles = [
    ('SUPER_ADMIN', 'National', 'All analytics, all governorates, user management, SIDRA export, terminology editing, raw data export, all Intelligence Labs, AuditLog review'),
    ('GLOBAL_ADMIN', 'National', 'Full analytics and all Labs. Cannot manage users or edit terminology.'),
    ('REGIONAL_ANALYST', 'Governorate', 'Own governorate data only, enriched with anonymized national benchmarks. Hard server-side scope enforcement overrides URL parameters.'),
    ('PRACTITIONER', 'School', 'School-level aggregates. Access to questionnaire form, class report workflow, QR distribution, practitioner guide. No geographic map.'),
    ('OPERATOR', 'School', 'Data entry only. Access to questionnaire submission, OCR scanning, session collection. No analytics.'),
]
make_table(doc, ['Role', 'Scope', 'Permissions'], roles, bg='0d6efd')

doc.add_paragraph()
add_p(doc, 'Security Implementation Detail: The HomepageView and SectionStatsView both implement hard authority limits. Even if a REGIONAL_ANALYST manipulates URL scope parameters to request national data, the server-side role check overrides the request and confines the response to the user\'s registered governorate. All login events, data access events, user management actions, and terminology changes are recorded in the AuditLog model with user, action, IP address, and timestamp.', italic=True)

# ------ SECTION 7 ------
doc.add_paragraph()
add_h(doc, '7. API Reference', 1)
endpoints = [
    ('Authentication', '/api/auth/', 'invite, activate, login, refresh, profile, pending-approvals, approve, reject, user list, user export CSV, delete user'),
    ('Geography', '/api/geography/', 'governorates list, establishments list (filterable by governorate)'),
    ('Questionnaire', '/api/questionnaire/', 'submit (atomic nested 21-section create), export (streaming CSV), submissions list, submission detail'),
    ('Statistics', '/api/stats/', 'school stats, governorate stats, national stats, insights, raw data export, regional profile'),
    ('Dashboard', '/api/', 'homepage (scoped KPIs + map + rankings), section-stats/<id>, lab-stats'),
    ('SIDRA', '/api/sidra/', 'export (national aggregated prevalence indicators)'),
    ('Class Reports', '/api/class-report/', 'create, list, detail, finalize, latest active'),
    ('Dynamic Questions', '/api/dynamic-questions/', 'list, detail/update by code'),
    ('Terminology', '/api/terminology/', 'list/create, update by key (SUPER_ADMIN, AuditLogged)'),
]
make_table(doc, ['Namespace', 'Base Path', 'Key Endpoints'], endpoints)

# ------ SECTION 8 ------
doc.add_paragraph()
add_h(doc, '8. Risk Register', 1)
risks = [
    ('RISK-01', 'HIGH', 'OCR misinterpretation of handwriting', 'Human-in-the-loop IntegrityLab validation UI; strict JSON schema enforcement; field-by-field review before submission'),
    ('RISK-02', 'HIGH', 'Student PII exposure via location data', 'Coordinates encrypted at rest; DBSCAN enforces minimum 5-point clusters; only centroids rendered; individual pins never shown'),
    ('RISK-03', 'HIGH', 'Cross-governorate data leakage', 'Server-side hard authority overrides in HomepageView and SectionStatsView; REGIONAL_ANALYST scope confined regardless of URL params'),
    ('RISK-04', 'MEDIUM', 'Correlation engine timeout on large datasets', 'Asynchronous Celery tasks + Redis; frontend polls every 2 seconds; 4-step progress state machine'),
    ('RISK-05', 'MEDIUM', 'Questionnaire schema drift from official MedSPAD paper', 'DynamicQuestion engine for admin-configurable extension; schema versioned via Django migrations'),
    ('RISK-06', 'MEDIUM', 'Brute-force authentication attacks', 'failed_attempts counter; SecureTokenObtainPairSerializer lockout; all login events AuditLogged with IP'),
    ('RISK-07', 'LOW', 'Database performance at scale', 'Computed boolean flags enable O(1) prevalence queries; select_related() used throughout; clear PostgreSQL migration path'),
]
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
header_row(t, ['ID', 'Severity', 'Risk Description', 'Mitigation Strategy'])
for risk in risks:
    row = t.add_row().cells
    for i, val in enumerate(risk):
        row[i].text = val
    if risk[1] == 'HIGH':
        set_cell_bg(row[1], 'f8d7da')
    elif risk[1] == 'MEDIUM':
        set_cell_bg(row[1], 'fff3cd')
    else:
        set_cell_bg(row[1], 'd1e7dd')

# ------ SECTION 9 ------
doc.add_paragraph()
add_h(doc, '9. Success Criteria', 1)
kpis = [
    ('SC-01', '100% structural parity with official MedSPAD 2026 questionnaire (21 sections, all choice scales)'),
    ('SC-02', 'Dashboard homepage load time < 2 seconds for all scoped aggregate queries'),
    ('SC-03', 'Zero PII exposures — all coordinate outputs anonymized via DBSCAN centroid clustering'),
    ('SC-04', 'OCR pipeline > 90% field extraction accuracy on clean scan input'),
    ('SC-05', 'Stress prediction Voting Regressor achieves R2 > 0.75 on held-out validation set'),
    ('SC-06', 'Platform operational across all 24 Tunisian governorates'),
    ('SC-07', '100% of data access events recorded in AuditLog with user, action, IP, and timestamp'),
    ('SC-08', 'SIDRA export endpoint returns valid national indicators within 5 seconds'),
]
make_table(doc, ['Criterion', 'Measurable Target'], kpis)

# ------ SECTION 10 ------
doc.add_paragraph()
add_h(doc, '10. Stakeholders', 1)
stakeholders = [
    ('Project Sponsor', 'Ministry of Health / ONMNE', 'Strategic oversight, budget approval, policy mandate'),
    ('Field Users', 'Practitioners, Operators (school staff)', 'Questionnaire collection, OCR scanning, class report management'),
    ('Analytical Users', 'Regional Analysts, Global Admins', 'Dashboard analytics, section exploration, regional profiling'),
    ('Platform Governance', 'Super Admins (national coordinators)', 'User management, system config, SIDRA federation, Intelligence Labs'),
    ('Data Consumers', 'SIDRA Platform, Public Health Researchers', 'Aggregated national indicators via /api/sidra/export/'),
    ('Development Team', 'Engineers, Data Scientists, UX Designers, Medical SMEs', 'Design, implementation, validation, deployment'),
]
make_table(doc, ['Stakeholder', 'Profile', 'Responsibility'], stakeholders)

doc.save('Project_PID_Sentinelle.docx')
print('PID saved successfully.')
